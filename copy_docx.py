# !/usr/bin/env python
# -*- coding:utf-8 -*-

"""
拷贝多个.docx文件指定内容到目标.docx文件中

"""

from Tkinter import *
import tkFileDialog
import docx
import os


def transition_context(file_path, destination_document):
    '''
    提取传入.docx文件有用信息，并追加到新.docx文件中
    '''
    doc_text = []
    num_paragraph = 0


    source_document = docx.Document(file_path)

    #doc_text = [paragraph.text for paragraph in source_document.paragraphs]

    for paragraph in source_document.paragraphs:
        if paragraph.text[0:8] == 'Keywords' or paragraph.text[0:8] == 'keywords' or paragraph.text[0:7] == 'keyword':
            break
        else:
            doc_text.append(paragraph.text)
            num_paragraph +=1
    c = file_path.split("\\")
    c = c[-1].split(".")
    doc_text[0] = doc_text[0] + "  " + c[0]
    

#   destination_document.add_heading('Document Test', 0)
    '''
    for i in range(num_paragraph):
        if (i >= 2) and (i%2==0):
            continue
        if i == num_paragraph-1:
            last_paragraph = destination_document.add_paragraph("")
            last_paragraph.add_run("Abstract: \n").bold = True
            last_paragraph.add_run(doc_text[i])
        else:
            destination_document.add_paragraph(doc_text[i])
    '''

    for i in range(num_paragraph):
        if (i >= 2) and (i%2==0):
            continue
        if (i == 0):
            new_paragraph = destination_document.add_paragraph(doc_text[i])
        elif(i == 1):
            new_paragraph.add_run("\n"  + doc_text[i])
        elif(i == num_paragraph-1):
            new_paragraph.add_run("\nAbstract: ").bold = True
            new_paragraph.add_run(doc_text[i])        
        else:
            new_paragraph.add_run(","  + doc_text[i])


def getdir(filepath=os.getcwd()):
    """
    用于获取目录下的文件列表
    """
    file_path = []
    files = os.listdir(filepath)
    for f in files:
        if f.lower().endswith('docx'):
            file_path.append(os.path.join(filepath, f))
    return file_path

def callback():
    entry.delete(0,END) #清空entry里面的内容
#   listbox_filename.delete(0,END)
    #调用filedialog模块的askdirectory()函数去打开文件夹
    
    filepath = tkFileDialog.askdirectory() 
    if filepath:
        entry.insert(0,filepath) #将选择好的路径加入到entry里面
    print (filepath)
    destination_document = docx.Document()
    for f_path in getdir(filepath):
        transition_context(f_path, destination_document)

    dit_file_path = filepath + "\\tango.docx"
    print dit_file_path
    destination_document.save(dit_file_path)
 
 
if __name__ == "__main__":
    root = Tk()
    root.title("Docx Tools")
    root.geometry("530x200")
    root.rowconfigure(1, weight=1)
    root.rowconfigure(2, weight=8)
            
    label1 = Label(root,text="Version: 0.02")   # 创建标签  
    label1.grid(sticky=W+N, row=1, column=1, padx=0, pady=0)

    label2 = Label(root,text="Author: tangaoo@126.com (if you find some bug, report me!)")   # 创建标签  
    label2.grid(sticky=W+N, row=2, column=1, padx=0, pady=0)

    entry = Entry(root, width=60)
    entry.grid(sticky=W+N, row=3, column=0, columnspan=4, padx=5, pady=20)
    
    button = Button(root,text="select folder",command=callback)
    button.grid(sticky=W+N, row=3, column=5, padx=5, pady=20)

    '''
    #创建loistbox用来显示所有文件名
    listbox_filename = Listbox(root, width=60)
    listbox_filename.grid(row=2, column=0, columnspan=4, rowspan=4, 
                            padx=5, pady=5, sticky=W+E+S+N)
    '''
    root.mainloop()